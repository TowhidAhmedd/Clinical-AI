"""
RAG pipeline: document loading, text extraction, chunking, and processing.
Supports PDF, DOCX, and TXT files.
"""
import re
import hashlib
from pathlib import Path
from typing import List, Optional
from pydantic import BaseModel
from loguru import logger

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    logger.warning("langchain-text-splitters not available — using simple chunking")


# --- Models ---
class DocumentChunk(BaseModel):
    chunk_id: str
    text: str
    metadata: dict
    embedding: Optional[List[float]] = None


class ProcessedDocument(BaseModel):
    doc_id: str
    filename: str
    total_chunks: int
    chunks: List[DocumentChunk]


# --- Text Extraction ---

def extract_text_from_pdf(file_path: str) -> List[dict]:
    """Extract text from PDF with page numbers."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"text": text, "page": i + 1})
        logger.info(f"Extracted {len(pages)} pages from {file_path}")
        return pages
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return []


def extract_text_from_docx(file_path: str) -> List[dict]:
    """Extract text from DOCX."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = []
        current_page = 1
        page_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                page_text.append(para.text.strip())
                # Simulate page breaks every 40 paragraphs
                if len(page_text) >= 40:
                    paragraphs.append({"text": "\n".join(page_text), "page": current_page})
                    current_page += 1
                    page_text = []
        if page_text:
            paragraphs.append({"text": "\n".join(page_text), "page": current_page})
        logger.info(f"Extracted {len(paragraphs)} sections from DOCX")
        return paragraphs
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return []


def extract_text_from_txt(file_path: str) -> List[dict]:
    """Extract text from TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        # Split into ~500 word chunks per "page"
        words = content.split()
        pages = []
        chunk_size = 500
        for i in range(0, len(words), chunk_size):
            page_words = words[i : i + chunk_size]
            pages.append({"text": " ".join(page_words), "page": (i // chunk_size) + 1})
        return pages
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return []


def extract_text(file_path: str, filename: str) -> List[dict]:
    """Route to the correct extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext == "docx":
        return extract_text_from_docx(file_path)
    elif ext == "txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# --- Text Cleaning ---

def clean_text(text: str) -> str:
    """Clean extracted text."""
    # Remove excessive whitespace
    text = re.sub(r"\s+", " ", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E\n]", "", text)
    # Remove repeated punctuation
    text = re.sub(r"\.{3,}", "...", text)
    return text.strip()


# --- Chunking ---

def chunk_text(
    pages: List[dict],
    doc_id: str,
    filename: str,
    chunk_size: int = 512,
    chunk_overlap: int = 64,
) -> List[DocumentChunk]:
    """
    Split pages into overlapping chunks using recursive character splitting.
    Each chunk retains metadata: source file, page number, chunk index.
    """
    if LANGCHAIN_AVAILABLE:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len,
        )
    else:
        # Simple fallback splitter
        class SimpleSplitter:
            def __init__(self, size, overlap):
                self.size = size
                self.overlap = overlap

            def split_text(self, text: str) -> List[str]:
                chunks = []
                start = 0
                while start < len(text):
                    end = min(start + self.size, len(text))
                    chunks.append(text[start:end])
                    start += self.size - self.overlap
                return [c for c in chunks if c.strip()]

        splitter = SimpleSplitter(chunk_size, chunk_overlap)

    chunks = []
    chunk_index = 0

    for page in pages:
        raw_text = clean_text(page["text"])
        page_num = page.get("page", 1)

        if not raw_text:
            continue

        splits = splitter.split_text(raw_text)

        for split_text in splits:
            if not split_text.strip():
                continue

            # Generate stable chunk ID
            chunk_hash = hashlib.md5(
                f"{doc_id}:{page_num}:{chunk_index}:{split_text[:50]}".encode()
            ).hexdigest()[:12]
            chunk_id = f"{doc_id[:8]}-p{page_num}-c{chunk_index}-{chunk_hash}"

            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    text=split_text,
                    metadata={
                        "doc_id": doc_id,
                        "filename": filename,
                        "page": page_num,
                        "chunk_index": chunk_index,
                        "chunk_id": chunk_id,
                        "text_length": len(split_text),
                    },
                )
            )
            chunk_index += 1

    logger.info(f"Created {len(chunks)} chunks from {filename}")
    return chunks


# --- Full Pipeline ---

def process_document(
    file_path: str,
    filename: str,
    doc_id: Optional[str] = None,
) -> ProcessedDocument:
    """
    Full processing pipeline:
    Upload → Extract → Clean → Chunk
    Returns ProcessedDocument with all chunks ready for embedding.
    """
    if doc_id is None:
        doc_id = hashlib.md5(f"{filename}:{Path(file_path).stat().st_size}".encode()).hexdigest()[:16]

    logger.info(f"Processing document: {filename} (doc_id={doc_id})")

    # 1. Extract text
    pages = extract_text(file_path, filename)
    if not pages:
        raise ValueError(f"No text could be extracted from {filename}")

    # 2. Chunk
    chunks = chunk_text(pages, doc_id=doc_id, filename=filename)
    if not chunks:
        raise ValueError(f"No chunks created from {filename}")

    return ProcessedDocument(
        doc_id=doc_id,
        filename=filename,
        total_chunks=len(chunks),
        chunks=chunks,
    )
